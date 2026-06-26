#!/usr/bin/env python3
"""Build the Twin KMS QA Test Plan as a Word .docx.

Consumes the screenshots captured by
`lightrag_webui_twin/e2e/capture-qa-screenshots.spec.ts`
(staged under docs/qa/assets/screenshots/) and renders a single,
self-contained outsourcing pack covering:

  - methodology (ID scheme, severity scale, roles, bug template)
  - known limitations testers must NOT report (audit caveats C0-C6)
  - the user-story catalogue
  - screen-by-screen, step-by-step test scripts (with screenshots)
  - the two outsourcing streams: E2E automation vendor + offshore manual QA
  - a traceability matrix

Run:  .venv/bin/python scripts/qa/build_qa_docx.py
Out:  docs/qa/Twin-KMS-QA-Test-Plan.docx
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
SHOTS = os.path.join(REPO, "docs", "qa", "assets", "screenshots")
OUT = os.path.join(REPO, "docs", "qa", "Twin-KMS-QA-Test-Plan.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
DANGER = RGBColor(0xB0, 0x2A, 0x2A)
GREEN = RGBColor(0x1F, 0x7A, 0x3D)

SEV_FILL = {"S1": "F4CCCC", "S2": "FCE5CD", "S3": "FFF2CC", "S4": "D9EAD3"}


# --------------------------------------------------------------------------- #
# low-level helpers
# --------------------------------------------------------------------------- #
def shade(cell, hex_fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcpr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def para(doc, text="", *, italic=False, bold=False, color=None, size=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.italic = italic
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
    return p


def bullet(doc, text, *, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def image(doc, filename, caption):
    path = os.path.join(SHOTS, filename)
    if not os.path.exists(path):
        para(doc, f"[missing screenshot: {filename}]", italic=True, color=DANGER)
        return
    doc.add_picture(path, width=Inches(6.4))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure — {caption}")
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    cap.paragraph_format.space_after = Pt(10)


def table(doc, headers, rows, *, widths=None, sev_col=None, style="Light Grid Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, bold=True, size=9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=9)
            if sev_col is not None and i == sev_col:
                fill = SEV_FILL.get(str(val).split()[0].strip())
                if fill:
                    shade(cells[i], fill)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def testcases(doc, cases):
    """cases: list of dict(id,title,type,sev,pre,steps,expected,note)."""
    for c in cases:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{c['id']} — {c['title']}")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = ACCENT
        # meta line
        m = doc.add_paragraph()
        m.paragraph_format.space_after = Pt(2)
        tr = m.add_run(f"Type: {c['type']}")
        tr.font.size = Pt(8.5)
        tr.font.color.rgb = MUTED
        sr = m.add_run(f"     Max severity if it fails: {c['sev']}")
        sr.font.size = Pt(8.5)
        sr.bold = True
        sr.font.color.rgb = DANGER if c["sev"] in ("S1", "S2") else MUTED
        if c.get("pre"):
            pp = doc.add_paragraph()
            pp.paragraph_format.space_after = Pt(2)
            lead = pp.add_run("Preconditions: ")
            lead.bold = True
            lead.font.size = Pt(9)
            body = pp.add_run(c["pre"])
            body.font.size = Pt(9)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)
        sp.add_run("Steps:").bold = True
        for step in c["steps"]:
            s = doc.add_paragraph(style="List Number")
            s.paragraph_format.space_after = Pt(1)
            s.add_run(step).font.size = Pt(9)
        ep = doc.add_paragraph()
        ep.paragraph_format.space_after = Pt(2)
        el = ep.add_run("Expected result: ")
        el.bold = True
        el.font.color.rgb = GREEN
        el.font.size = Pt(9)
        eb = ep.add_run(c["expected"])
        eb.font.size = Pt(9)
        if c.get("note"):
            npar = doc.add_paragraph()
            npar.paragraph_format.space_after = Pt(6)
            nl = npar.add_run("Note: ")
            nl.bold = True
            nl.font.size = Pt(8.5)
            nb = npar.add_run(c["note"])
            nb.font.size = Pt(8.5)
            nb.italic = True


def callout(doc, title, lines, fill="FFF2CC"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    cell.text = ""
    tp = cell.paragraphs[0]
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(10)
    for ln in lines:
        lp = cell.add_paragraph()
        lp.paragraph_format.space_after = Pt(2)
        lp.add_run("• ").bold = True
        lp.add_run(ln).font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(el)


def page_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Twin KMS — QA Test Plan & Outsourcing Pack · CONFIDENTIAL · v1.0 (2026-06-26)")
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED


# --------------------------------------------------------------------------- #
# document
# --------------------------------------------------------------------------- #
def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ---- cover -----------------------------------------------------------
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Twin KMS")
    r.bold = True
    r.font.size = Pt(40)
    r.font.color.rgb = ACCENT
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("QA Test Plan & Outsourcing Pack")
    rs.font.size = Pt(20)
    rs.font.color.rgb = MUTED
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsub = sub.add_run("Operator console — user stories, user flows, and screen-by-screen test scripts\n"
                       "for E2E automation and manual (offshore) testing")
    rsub.font.size = Pt(12)
    for _ in range(8):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run("Version 1.0  ·  2026-06-26  ·  CONFIDENTIAL\n"
                      "Scope: lightrag_webui_twin (React 19 operator console) + /twin/api overlay")
    rm.font.size = Pt(11)
    rm.font.color.rgb = MUTED
    doc.add_page_break()

    # ---- TOC -------------------------------------------------------------
    h(doc, "Table of Contents", 1)
    add_toc(doc)
    doc.add_page_break()

    # ===================================================================== #
    # 1. Purpose & how to use
    # ===================================================================== #
    h(doc, "1. Purpose & how to use this document", 1)
    para(doc,
         "This document is the single source of truth for outsourced QA of the Twin KMS operator "
         "console. It maps the product's user stories and user flows to concrete, executable test "
         "cases, and splits them across two delivery streams:")
    bullet(doc, "Automated end-to-end (E2E) regression of operator journeys against a deterministic, "
                "mocked backend (and a real backend variant). Owned by the test-automation vendor.",
           bold_lead="Stream A — E2E automation. ")
    bullet(doc, "Human, exploratory and judgement-based testing: visual fidelity, real-document "
                "ingestion, answer quality, cross-browser behaviour, and role-based access. Owned by "
                "the offshore manual QA team.",
           bold_lead="Stream B — Manual (offshore). ")
    para(doc,
         "Each test case below is tagged with its stream and the maximum severity to assign if it "
         "fails. The automation vendor implements / extends the E2E suite; the manual team executes "
         "the Manual cases and reports defects using the template in §3.5.")
    callout(doc, "Read §4 (Known limitations) BEFORE testing.",
            ["Several behaviours look like bugs but are documented, accepted product limitations of "
             "the 1.0.0 release. Reporting them wastes triage time. They are listed in §4 with the "
             "audit reference."],
            fill="FCE5CD")

    # ===================================================================== #
    # 2. Test environments
    # ===================================================================== #
    h(doc, "2. Test environments & access", 1)
    para(doc, "Two builds exist. Choose the build that matches the test's intent.")
    h(doc, "2.1 Mocked build (MSW) — default for UI / flow / visual testing", 2)
    para(doc,
         "The console ships with an in-browser mock backend (Mock Service Worker). All screens render "
         "realistic, deterministic data with no database, no LLM and no network dependency. This is the "
         "correct environment for Stream A regression and for most Stream B screen/flow testing.")
    bullet(doc, "cd lightrag_webui_twin && bun install && bun run dev", bold_lead="Boot: ")
    bullet(doc, "The dev user is an administrator, so every admin-gated control is visible.", bold_lead="Identity: ")
    bullet(doc, "LLM answers are mock strings (\"Mock retrieval response for: …\") and retrieval sources "
                "are fixtures. Answer quality cannot be judged here.", bold_lead="Caveat: ")
    h(doc, "2.2 Real-backend build — for ingestion & retrieval-quality testing", 2)
    para(doc,
         "A full runtime (Memgraph + LightRAG + LLM credentials) is required to test document ingestion, "
         "classification, real retrieval/answer quality and folder data isolation. The integration "
         "environment URL and credentials are provided separately by the Twincore team. Stream B uses "
         "this build only for the cases explicitly marked \"Real backend\".")
    h(doc, "2.3 Roles & permissions", 2)
    para(doc, "Access is tiered. Many controls are hidden or rejected based on the operator's role. "
              "Stream B must verify gating with each role where a real/IdP-configured environment is available.")
    table(doc,
          ["Role / tier", "Can do", "Cannot do"],
          [
              ["Reader (Tier 1)", "Read documents, tags, graph, activity; run retrieval queries.",
               "Approve/reject docs; create/edit/delete tags; mutate graph; manage folders or API keys."],
              ["Steward / Contributor (Tier 2)", "Everything a Reader can, plus approve/reject documents and "
               "suggest tag edits (queued for review).",
               "In-place tag edit/delete; folder CRUD (unless granted admin:folders)."],
              ["Admin (Tier 3)", "Full CRUD on tags, graph entities/relations, folders (with admin:folders "
               "scope), API keys, tag categories; clear activity.",
               "—"],
          ],
          widths=[1.6, 2.6, 2.2])
    para(doc,
         "Folder administration specifically requires the admin:folders gateway scope. In a mocked build "
         "the dev user holds it; in a real build it is granted when the user's IdP groups intersect the "
         "configured admin groups.", italic=True, size=9)

    # ===================================================================== #
    # 3. Methodology
    # ===================================================================== #
    h(doc, "3. Methodology", 1)
    h(doc, "3.1 Identifier scheme", 2)
    table(doc,
          ["Prefix", "Meaning", "Example"],
          [
              ["US-xx", "User story (epic-grouped)", "US-DOC-03"],
              ["TC-xxx", "Test case (executed by a tester)", "TC-DOC-12"],
              ["E2E", "Already covered by an automated Playwright spec", "documents.spec.ts"],
          ],
          widths=[1.3, 3.4, 1.7])
    h(doc, "3.2 Test case types", 2)
    bullet(doc, "the journey is already automated; the manual team performs a one-pass smoke confirmation "
                "only, and the automation vendor maintains the spec.", bold_lead="E2E (automated): ")
    bullet(doc, "an automated journey that does not yet exist and is recommended for the vendor to build.",
           bold_lead="E2E (to build): ")
    bullet(doc, "requires human judgement (visual, semantic, real data, cross-browser) — owned by the "
                "offshore team.", bold_lead="Manual: ")
    h(doc, "3.3 Severity scale", 2)
    para(doc, "Use this scale for every defect. It mirrors the project's internal audit scale "
              "(bloquant / sérieux / modéré / info).")
    table(doc,
          ["Severity", "Label", "Definition", "Internal equivalent"],
          [
              ["S1", "Blocker", "A core operator flow is impossible; data loss/corruption; security or "
               "access-control bypass. Stop-ship.", "bloquant"],
              ["S2", "Critical", "A major feature is broken with no reasonable workaround, or wrong data is "
               "presented as authoritative.", "sérieux"],
              ["S3", "Major", "Feature degraded but a workaround exists; significant UX or layout defect.", "modéré"],
              ["S4", "Minor", "Cosmetic, copy, or minor visual issue; no functional impact.", "info"],
          ],
          widths=[0.8, 1.0, 3.3, 1.3], sev_col=0)
    h(doc, "3.4 Definition of done for a test pass", 2)
    bullet(doc, "Every Manual TC executed on the agreed browser matrix (§7.2) with a recorded verdict.")
    bullet(doc, "Each defect has: screenshot/video, exact steps, expected vs actual, severity, environment.")
    bullet(doc, "No defect filed for any item listed in §4 (Known limitations).")
    h(doc, "3.5 Defect report template (Stream B)", 2)
    table(doc,
          ["Field", "Content"],
          [
              ["ID", "BUG-<area>-<n> (e.g. BUG-DOC-07)"],
              ["Title", "One line: screen + symptom"],
              ["Screen / TC", "e.g. Documents / TC-DOC-12"],
              ["Environment", "Mocked or Real backend; browser + version; OS; viewport"],
              ["Role", "Reader / Steward / Admin"],
              ["Preconditions", "State required before step 1"],
              ["Steps", "Numbered, reproducible"],
              ["Expected", "What should happen (cite the TC)"],
              ["Actual", "What happened"],
              ["Severity", "S1–S4 per §3.3"],
              ["Evidence", "Screenshot / screen recording attached"],
          ],
          widths=[1.4, 5.0])

    # ===================================================================== #
    # 4. Known limitations — DO NOT REPORT
    # ===================================================================== #
    doc.add_page_break()
    h(doc, "4. Known limitations — DO NOT report these as defects", 1)
    para(doc,
         "The following behaviours are documented, accepted limitations of the 1.0.0 release, traced to "
         "the internal interaction audit (2026-06-13). They are expected. Do not raise defects for them. "
         "If a test step appears to contradict one of these, follow the limitation, not your intuition.")
    table(doc,
          ["#", "Behaviour you will observe", "Why it is expected", "Ref"],
          [
              ["1", "On the Retrieval tab, adding a Source-tag filter does NOT narrow the answer. The "
               "answer can still cite material outside the selected tags.",
               "LightRAG 1.4.9.11 has no tag_filter at retrieval time; the /query path rejects it. The "
               "filter is only honoured on the structured /query/data route.", "C1"],
              ["2", "The Retrieval \"Sources\" panel may list chunks that are not exactly what the model "
               "grounded its answer on.",
               "Sources are reconstructed by a secondary retrieval, not yet the model's own references. "
               "Treat them as candidate sources, not citations.", "C0/C3"],
              ["3", "The Knowledge Graph may show demo entities even when the real KB is empty.",
               "A seed fallback keeps the demo/pre-ingestion path renderable. In the mocked build ALL "
               "graph data is fixture data by design.", "C5"],
              ["4", "\"Read source\" / indexed-chunks preview may show placeholder text rather than the "
               "real extracted document text.",
               "Real extracted-text view is a later backend phase; the modal is a contract placeholder.", "C6"],
              ["5", "\"View raw\" in the document detail panel opens a notice instead of downloading.",
               "The raw-download endpoint is backend phase 2; the button is intentionally a notice gate.", "—"],
              ["6", "Folders filter the WebUI governance surface but do not create separate KB instances; "
               "graph reads are membership-scoped, not physically isolated.",
               "By design — one physical namespace, folders are a logical many-to-many membership layer.", "C4"],
          ],
          widths=[0.3, 2.6, 2.9, 0.5])

    # ===================================================================== #
    # 5. User-story catalogue
    # ===================================================================== #
    doc.add_page_break()
    h(doc, "5. User-story catalogue", 1)
    para(doc, "Stories are grouped by epic. Each maps to the screen section (§6) where its flows are tested.")
    table(doc,
          ["US", "As a… I want… so that…", "Tested in"],
          [
              ["US-AUTH-01", "operator, I sign in with my credentials so that only authorised staff reach the KB.", "§6.1"],
              ["US-AUTH-02", "operator, I sign out and have my local session/threads cleared.", "§6.1 / §6.10"],
              ["US-NAV-01", "operator, I switch between the active folders I am entitled to.", "§6.2"],
              ["US-NAV-02", "operator, I see and triage my notifications from any screen.", "§6.2"],
              ["US-DOC-01", "operator, I browse, search and filter the document inventory by status, tag and source.", "§6.3"],
              ["US-DOC-02", "operator, I inspect a document's chunks, lineage and audit trail.", "§6.3"],
              ["US-DOC-03", "steward, I approve or reject newly submitted / modified sources.", "§6.3"],
              ["US-DOC-04", "operator, I retag one or many documents and see the impact before applying.", "§6.3"],
              ["US-DOC-05", "operator, I delete a document (single or bulk) with a confirmation guard and undo.", "§6.3"],
              ["US-DOC-06", "operator, I reprocess failed documents.", "§6.3"],
              ["US-UPL-01", "operator, I upload files (or add URLs), assign tags, and set a sensitivity class.", "§6.4"],
              ["US-UPL-02", "operator, files above the allowed sensitivity ceiling are refused, not indexed.", "§6.4"],
              ["US-TAG-01", "operator, I request a new tag for review.", "§6.5"],
              ["US-TAG-02", "admin, I edit, deprecate, manage synonyms of, or delete a tag (with migration).", "§6.5"],
              ["US-TAG-03", "admin, I approve / reject pending tag requests.", "§6.5"],
              ["US-TAG-04", "admin, I manage the tag-category taxonomy (add, import, template).", "§6.5"],
              ["US-RET-01", "operator, I ask a question and receive an answer with a sources panel.", "§6.6"],
              ["US-RET-02", "operator, I tune retrieval parameters (mode, top-k, rerank, filters).", "§6.6"],
              ["US-RET-03", "operator, my conversation threads persist locally across reloads.", "§6.6"],
              ["US-GRAPH-01", "operator, I explore the knowledge graph and filter by type / tag / document.", "§6.7"],
              ["US-GRAPH-02", "operator, I inspect an entity and its relations and pin it.", "§6.7"],
              ["US-GRAPH-03", "admin, I create / edit / delete entities and relations.", "§6.7"],
              ["US-ACT-01", "operator, I review the audit feed and filter by time, kind, severity, actor.", "§6.8"],
              ["US-ACT-02", "operator, I open an event and drill down to its source document.", "§6.8"],
              ["US-SET-01", "operator, I view my profile, identity provider and gateway scopes.", "§6.9"],
              ["US-SET-02", "operator, I browse the live API surface.", "§6.9"],
              ["US-SET-03", "admin, I mint, reveal-once, and revoke per-operator API keys.", "§6.9"],
              ["US-SET-04", "admin, I create / edit / delete runtime folders within the cap.", "§6.9"],
          ],
          widths=[1.1, 4.4, 0.9])

    # ===================================================================== #
    # 6. Screen-by-screen test scripts
    # ===================================================================== #
    doc.add_page_break()
    h(doc, "6. Screen-by-screen test scripts", 1)
    para(doc, "Each subsection: a screenshot, what the screen does, the stories it covers, then the "
              "test cases. Cases marked E2E (automated) already have Playwright coverage — the manual "
              "team runs them once as a smoke check.")

    for screen in SCREENS:
        h(doc, screen["title"], 2)
        if screen.get("shot"):
            image(doc, screen["shot"], screen["caption"])
        para(doc, screen["overview"])
        for fn, cap in screen.get("extra_shots", []):
            image(doc, fn, cap)
        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(4)
        a = meta.add_run("Reach: ")
        a.bold = True
        a.font.size = Pt(9)
        meta.add_run(screen["reach"]).font.size = Pt(9)
        st = doc.add_paragraph()
        st.paragraph_format.space_after = Pt(6)
        b = st.add_run("Stories: ")
        b.bold = True
        b.font.size = Pt(9)
        st.add_run(", ".join(screen["stories"])).font.size = Pt(9)
        testcases(doc, screen["cases"])
        doc.add_page_break()

    # ===================================================================== #
    # 7. Stream scopes
    # ===================================================================== #
    h(doc, "7. Outsourcing scope by stream", 1)
    h(doc, "7.1 Stream A — E2E automation vendor", 2)
    para(doc,
         "The console already carries a substantial Playwright suite (18 specs) plus ~45 component "
         "tests. The vendor's job is to MAINTAIN these and CLOSE the gaps below — not to re-cover what "
         "exists. Existing coverage by area:")
    table(doc,
          ["Area", "Covered", "Spec(s)"],
          [
              ["Login / auth gate / session", "Yes", "login.spec.ts"],
              ["Documents CRUD, filter, bulk, pending review", "Yes", "documents.spec.ts, app.spec.ts"],
              ["Upload + classification header", "Yes", "upload.spec.ts"],
              ["Tag governance (request→approve→edit→migrate)", "Yes", "tags.spec.ts"],
              ["Retrieval threads, params, filters", "Partial", "retrieval.spec.ts (citation link skipped)"],
              ["Knowledge graph CRUD, pin, cascade", "Yes", "graph.spec.ts"],
              ["Activity filters & detail drill-down", "Yes", "activity.spec.ts"],
              ["Settings: API keys, profile, folders", "Yes", "settings.spec.ts"],
              ["Notifications, topbar, theme, responsive", "Yes", "toasts-topbar / responsive-topbar"],
              ["Modal accessibility (focus, escape)", "Yes", "modals-a11y.spec.ts"],
              ["Runtime folders & X-Twin-Folder header", "Yes", "folders-runtime.spec.ts"],
              ["Adversarial API surface (real backend)", "Yes", "api-coverage-real / -generated-key"],
          ],
          widths=[3.0, 1.0, 2.4])
    para(doc, "Recently closed (2026-06-26 reinforcement pass):", bold=True)
    bullet(doc, "Citation drilldown — a Retrieval source now links to its real seeded document; the "
                "previously-skipped test is enabled (retrieval.spec.ts, with a one-source MSW handler tune).")
    bullet(doc, "Role gating on folder administration — Reader (no admin:folders) sees a read-only badge "
                "and no Add-folder control; Admin sees the controls (role-gating.spec.ts).")
    bullet(doc, "Document pagination boundary — 60 seeded docs force a second page; Next/Previous and the "
                "page label are asserted (documents-pagination.spec.ts).")
    para(doc, "Gaps still recommended for new automation (E2E to build):", bold=True)
    bullet(doc, "Real-backend ingestion: upload a real file → poll to COMPLETED → appears with chunks.")
    bullet(doc, "Classification rejection end-to-end: a file above the ceiling lands FAILED + emits a "
                "classification-rejected activity event (MSW needs a classification-rejected seed knob).")
    bullet(doc, "Folder membership: duplicate-to-share, add-to-folder, and last-folder removal physically "
                "deletes the document (endpoints are mocked; an admin UI path / contract test is needed).")
    bullet(doc, "Quota-exceeded journey (QuotaBanner) — needs a /__e2e/quota knob to drive the over-limit state.")
    bullet(doc, "Session-expiry re-auth (real backend / short token TTL).")
    callout(doc, "Finding — Tags governance is NOT role-gated in the UI.",
            ["TagsTab is wired to a hardcoded CURRENT_USER constant (palier 3 = admin/steward), not to the "
             "authenticated user. A Reader therefore still sees Edit/Delete/Deprecate tag controls. The MSW "
             "tag-mutation handlers are also not scope-checked. Net: tag-governance access control is "
             "currently enforced (if at all) only by the real backend, not the UI. Decide: wire CURRENT_USER "
             "→ auth.user (so the UI reflects the role), or accept backend-only enforcement and document it. "
             "Until resolved, a UI role-gating test for tags would be vacuous — which is why role-gating.spec.ts "
             "covers folders only."],
            fill="FCE5CD")

    h(doc, "7.2 Stream B — manual (offshore) testing", 2)
    para(doc, "Human judgement work that automation cannot meaningfully cover.")
    para(doc, "Browser / device matrix (minimum):", bold=True)
    table(doc,
          ["Browser", "Versions", "Viewports"],
          [
              ["Chrome", "latest + latest-1", "1440×900, 1280×800"],
              ["Firefox", "latest", "1440×900"],
              ["Safari", "latest (macOS)", "1440×900"],
              ["Edge", "latest", "1280×800"],
              ["Responsive", "Chrome", "1100px and 760px widths"],
          ],
          widths=[1.6, 2.2, 2.6])
    para(doc, "Exploratory charters (time-boxed, 60–90 min each):", bold=True)
    bullet(doc, "Real-document ingestion: ingest 10 mixed real files (PDF, DOCX, XLSX, HTML) and verify "
                "they reach COMPLETED with sensible chunk counts and tags.")
    bullet(doc, "Answer quality (real backend): ask 15 domain questions; judge whether the answer is "
                "grounded, whether sources are relevant, and whether a deliberately out-of-scope question "
                "honestly returns \"insufficient information\".")
    bullet(doc, "Classification: upload files carrying MIP labels C1–C4 and confirm the badge shown and "
                "the ceiling-rejection behaviour.")
    bullet(doc, "Visual fidelity & dark mode: every screen in light and dark theme; check contrast, "
                "truncation, overflow, and long filenames/tags.")
    bullet(doc, "Role-based access: repeat the golden paths as Reader, Steward, Admin; confirm gated "
                "controls are correctly hidden/blocked.")
    bullet(doc, "Resilience: slow network, mid-flight reloads, double-clicks on destructive actions, "
                "empty states, and error toasts (undo).")

    # ===================================================================== #
    # 8. Traceability matrix
    # ===================================================================== #
    doc.add_page_break()
    h(doc, "8. Traceability matrix (story → screen → cases → automation)", 1)
    table(doc,
          ["User story", "Screen", "Key test cases", "Automated?"],
          [
              ["US-AUTH-01/02", "§6.1 Login", "TC-AUTH-01..04", "Yes (login.spec)"],
              ["US-NAV-01/02", "§6.2 Topbar", "TC-NAV-01..05", "Yes (folders-runtime, toasts-topbar)"],
              ["US-DOC-01..06", "§6.3 Documents", "TC-DOC-01..12", "Mostly (documents.spec)"],
              ["US-UPL-01/02", "§6.4 Upload", "TC-UPL-01..06", "Partial (upload.spec; rejection manual)"],
              ["US-TAG-01..04", "§6.5 Tags", "TC-TAG-01..08", "Yes (tags.spec)"],
              ["US-RET-01..03", "§6.6 Retrieval", "TC-RET-01..07", "Partial (citation/quality manual)"],
              ["US-GRAPH-01..03", "§6.7 Graph", "TC-GR-01..08", "Yes (graph.spec)"],
              ["US-ACT-01/02", "§6.8 Activity", "TC-ACT-01..06", "Yes (activity.spec)"],
              ["US-SET-01..04", "§6.9 Settings", "TC-SET-01..09", "Yes (settings.spec)"],
          ],
          widths=[1.4, 1.4, 1.9, 1.8])
    para(doc, "End of document.", italic=True, color=MUTED)

    page_footer(doc)
    doc.save(OUT)
    print(f"wrote {OUT}")


# --------------------------------------------------------------------------- #
# screen content
# --------------------------------------------------------------------------- #
SCREENS = [
    {
        "title": "6.1 Login & session",
        "shot": "01-login.png",
        "caption": "Auth gate (shown when an authentication backend is configured).",
        "overview": "The login gate appears only when the deployment configures an auth backend "
                    "(API key, local JWT, or IdP). In the open-access default it is skipped. Local login "
                    "issues an HttpOnly session cookie; sign-out clears it and purges local conversation threads.",
        "reach": "First screen on load when auth is enabled; sign-out from Settings › Profile returns here.",
        "stories": ["US-AUTH-01", "US-AUTH-02"],
        "cases": [
            {"id": "TC-AUTH-01", "title": "Submit is disabled until both fields are filled",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Auth gate enabled; login screen shown.",
             "steps": ["Observe the Sign-in button is disabled.",
                       "Type a username only — button stays disabled.",
                       "Type a password — button becomes enabled."],
             "expected": "Sign-in enables only when both username and password are non-empty; the app "
                         "shell is not visible behind the gate."},
            {"id": "TC-AUTH-02", "title": "Bad credentials show an error and keep the gate",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Login screen shown.",
             "steps": ["Enter a valid username and a wrong password.", "Click Sign in."],
             "expected": "An inline error appears; the login screen remains; no app content leaks."},
            {"id": "TC-AUTH-03", "title": "Valid credentials enter the app and survive reload",
             "type": "E2E (automated)", "sev": "S1",
             "pre": "Login screen shown.",
             "steps": ["Enter valid credentials and click Sign in.",
                       "Confirm the Documents screen loads.",
                       "Reload the page."],
             "expected": "The session persists across reload; the operator stays signed in."},
            {"id": "TC-AUTH-04", "title": "Sign-out returns to the gate and clears local threads",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Signed in.",
             "steps": ["Go to Settings › Profile.", "Click Sign out.",
                       "Inspect browser localStorage for twin-rag.* keys."],
             "expected": "The login gate returns and all twin-rag.* local keys are removed."},
            {"id": "TC-AUTH-05", "title": "Session expiry forces re-authentication",
             "type": "Manual", "sev": "S2",
             "pre": "Real backend with a short token TTL, signed in.",
             "steps": ["Stay idle until the token expires.",
                       "Trigger any authenticated action (e.g. open Documents)."],
             "expected": "The operator is challenged to re-authenticate; no stale data is shown as live.",
             "note": "Real-backend only; the mocked build does not expire tokens."},
        ],
    },
    {
        "title": "6.2 Topbar — navigation, folders, notifications",
        "shot": "21-folder-switcher.png",
        "caption": "Folder switcher menu open in the topbar.",
        "overview": "The persistent topbar carries: brand (returns to Documents), the six tab buttons, the "
                    "folder switcher, the notifications bell with an unread badge, and the theme toggle. The "
                    "folder switcher lists only folders the operator is entitled to; switching propagates an "
                    "X-Twin-Folder header on subsequent requests and resets document filters.",
        "reach": "Always visible once authenticated.",
        "extra_shots": [("20-notifications.png", "Notifications popover with unread items and quick actions.")],
        "stories": ["US-NAV-01", "US-NAV-02"],
        "cases": [
            {"id": "TC-NAV-01", "title": "Tab navigation reaches all six tabs",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Signed in.",
             "steps": ["Click each tab: Documents, Tags, Retrieval, Graph, Activity, Settings."],
             "expected": "Each tab renders its primary heading with no console errors."},
            {"id": "TC-NAV-02", "title": "Folder switch changes scope and resets filters",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "More than one folder available.",
             "steps": ["Open the folder switcher.", "Select a different folder.",
                       "Observe the Documents list and any active filters."],
             "expected": "The active folder updates, the document list refreshes for the new folder, and "
                         "local filters are cleared."},
            {"id": "TC-NAV-03", "title": "Empty-folder state shows the Twincore guidance",
             "type": "Manual", "sev": "S3",
             "pre": "A deployment with no folder provisioned.",
             "steps": ["Load the console with no folder available."],
             "expected": "The message \"No folder available for this KB. Please contact Twincore Team.\" is shown."},
            {"id": "TC-NAV-04", "title": "Notifications popover: mark-all-read and clear",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "At least one unread notification.",
             "steps": ["Click the bell.", "Click Mark all read.", "Re-open and click Clear all."],
             "expected": "The unread badge clears after mark-all-read; the list empties after clear; "
                         "\"View full activity log\" navigates to Activity."},
            {"id": "TC-NAV-05", "title": "Theme toggle persists across reload",
             "type": "Manual", "sev": "S4",
             "pre": "Signed in.",
             "steps": ["Toggle to dark mode.", "Reload."],
             "expected": "Dark mode persists; every screen remains legible with adequate contrast.",
             "note": "Pair with the dark-mode visual charter in §7.2."},
        ],
    },
    {
        "title": "6.3 Documents",
        "shot": "02-documents-list.png",
        "caption": "Documents tab: pending-review cards above the filterable inventory table.",
        "overview": "The operator's main workspace. A pending-review band surfaces newly submitted or "
                    "upstream-modified sources for approval; below it, a filterable, paginated inventory with "
                    "status, classification badge, tags, chunk count and per-row actions. A pipeline popover "
                    "shows ingestion status and a Retry-failed control.",
        "reach": "Default tab after login; brand button returns here.",
        "extra_shots": [
            ("04-document-detail.png", "Document detail panel: chunks, lineage and audit tabs."),
            ("05-retag-modal.png", "Retag modal with current tags and an impact preview."),
        ],
        "stories": ["US-DOC-01", "US-DOC-02", "US-DOC-03", "US-DOC-04", "US-DOC-05", "US-DOC-06"],
        "cases": [
            {"id": "TC-DOC-01", "title": "Search and status filter narrow the list",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Several documents across statuses.",
             "steps": ["Type a filename fragment in search.",
                       "Click the Failed status pill.", "Clear filters."],
             "expected": "The table shows only matching rows; status pills reflect backend counts; clearing restores all."},
            {"id": "TC-DOC-02", "title": "Tag and source filters combine",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Documents carrying tags.",
             "steps": ["Open the tag filter, add a tag.", "Add a source filter."],
             "expected": "Only documents matching all active filters remain; filter chips are removable."},
            {"id": "TC-DOC-03", "title": "Open the document detail panel",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "At least one document.",
             "steps": ["Click a document row.", "Inspect the Chunks, Lineage and Audit tabs."],
             "expected": "The side panel shows chunk content, metadata (created/updated, folder, tags, "
                         "classification) and related audit events."},
            {"id": "TC-DOC-04", "title": "Approve a pending (newly submitted) source",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "A card in the \"Requested\" pending state; Steward+ role.",
             "steps": ["Click Read source to preview indexed chunks.", "Click Approve."],
             "expected": "The source leaves the pending band; an approval activity event is recorded; "
                         "focus returns to a sensible element."},
            {"id": "TC-DOC-05", "title": "Reject a pending source with a reason",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "A pending card; Steward+ role.",
             "steps": ["Click Reject.", "Confirm the reject button is disabled until a reason is typed.",
                       "Enter a reason and submit."],
             "expected": "Rejection requires a reason; the card is removed; a rejection event is recorded."},
            {"id": "TC-DOC-06", "title": "Re-validate an upstream-modified source",
             "type": "Manual", "sev": "S2",
             "pre": "A card in the \"Modified\" pending state.",
             "steps": ["Open the modified card.", "Review the change summary.", "Click Approve update or Reject update."],
             "expected": "The modified source is re-validated or rejected; the band updates accordingly."},
            {"id": "TC-DOC-07", "title": "Single-document retag with impact preview",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "A document open or row action available.",
             "steps": ["Open Retag for one document.", "Add a tag and remove one.",
                       "Read the impact preview (chunks/docs affected).", "Apply."],
             "expected": "The preview reflects the change; on Apply the tags update optimistically and an "
                         "undo toast appears."},
            {"id": "TC-DOC-08", "title": "Bulk retag across multiple selected documents",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Two or more documents selectable.",
             "steps": ["Select multiple rows.", "Click Retag.", "Note shared vs partial tags.",
                       "Add a tag and Apply."],
             "expected": "Shared and partial (\"on some docs\") tags are distinguished; the change applies to "
                         "all selected documents."},
            {"id": "TC-DOC-09", "title": "Bulk retag failure rolls back optimistic tags",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Inject a 500/413 on bulk retag (MSW scenario).",
             "steps": ["Bulk-retag a selection.", "Observe the failure handling."],
             "expected": "The optimistic tags roll back and an error is surfaced; no partial/incorrect "
                         "state persists."},
            {"id": "TC-DOC-10", "title": "Delete (single and bulk) with confirm and undo",
             "type": "E2E (automated)", "sev": "S1",
             "pre": "Deletable documents.",
             "steps": ["Delete a single document via the detail panel (two-click confirm).",
                       "Select several and use bulk delete (double-confirm).",
                       "Observe the deleting state and the undo toast."],
             "expected": "Deletion requires explicit confirmation; rows disappear; an undo path is offered; "
                         "no accidental single-click deletion is possible."},
            {"id": "TC-DOC-11", "title": "Reprocess failed documents",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "At least one FAILED document.",
             "steps": ["Use Retry Failed (or Re-process in the detail panel on a FAILED doc).",
                       "Try Re-process on a non-failed doc."],
             "expected": "Failed documents are re-queued; a non-failed doc reports honestly that reprocess "
                         "does not apply."},
            {"id": "TC-DOC-12", "title": "Cascade: deleting a document updates the graph",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "A document whose entities appear on the Graph tab.",
             "steps": ["Delete the document.", "Open the Graph tab."],
             "expected": "Entities sourced only from the deleted document are no longer shown; the graph "
                         "refetches rather than serving stale nodes."},
        ],
    },
    {
        "title": "6.4 Upload & classification",
        "shot": "03-add-source-modal.png",
        "caption": "Add-source modal: file drop, tag assignment, and per-file sensitivity class.",
        "overview": "Sources are added by file upload (drag-drop or picker, ≤50 MB) or by URL "
                    "(Confluence/SharePoint). Each file can be tagged and assigned a MIP sensitivity class "
                    "(C1 Public → C4 Secret). Files above the workspace ceiling are refused at ingestion and "
                    "land FAILED rather than being indexed.",
        "reach": "\"Add source\" button on the Documents tab.",
        "stories": ["US-UPL-01", "US-UPL-02"],
        "cases": [
            {"id": "TC-UPL-01", "title": "Add a valid file and submit",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Add-source modal open.",
             "steps": ["Select a supported file under 50 MB.", "Wait for it to validate.", "Click Add 1 source."],
             "expected": "The file validates, the submit button enables, and the source is accepted for ingestion."},
            {"id": "TC-UPL-02", "title": "File type / size validation",
             "type": "Manual", "sev": "S3",
             "pre": "Add-source modal open.",
             "steps": ["Add an unsupported type.", "Add a file over 50 MB."],
             "expected": "Each invalid file shows a clear per-file error and cannot be submitted; valid files still submit."},
            {"id": "TC-UPL-03", "title": "Multi-file: only valid files submit",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Add-source modal open.",
             "steps": ["Add a mix of one valid and one invalid file.", "Submit."],
             "expected": "Counts are accurate; only the valid file is submitted; the failure is reported truthfully."},
            {"id": "TC-UPL-04", "title": "Assign initial tags on upload",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Add-source modal open.",
             "steps": ["Add a file.", "Search and select a tag from the catalog.", "Submit."],
             "expected": "The uploaded source carries the chosen tag once ingested (with manual-retag guidance "
                         "if tag propagation times out)."},
            {"id": "TC-UPL-05", "title": "Set a per-file sensitivity class",
             "type": "Manual", "sev": "S3",
             "pre": "Add-source modal open; real backend.",
             "steps": ["Add a file.", "Pick a class (e.g. C2 · Internal).", "Submit and inspect the document badge."],
             "expected": "The classification badge on the resulting document matches the chosen class (or the "
                         "label detected from the file)."},
            {"id": "TC-UPL-06", "title": "Over-ceiling file is rejected, not indexed",
             "type": "Manual", "sev": "S1",
             "pre": "Real backend with TWIN_MIP_MAX_CLASSIFICATION set (e.g. C2); a C3/C4-labelled file.",
             "steps": ["Upload the over-ceiling file.", "Open Documents.", "Open Activity."],
             "expected": "The document appears FAILED with a classification-rejected reason, is never indexed "
                         "(no chunks, not retrievable), and a classification-rejected warning event is logged.",
             "note": "This is the core compliance guarantee — verify it cannot be bypassed."},
        ],
    },
    {
        "title": "6.5 Tags governance",
        "shot": "06-tags-list.png",
        "caption": "Tags tab: category rail, status filters, tag cards and detail panel.",
        "overview": "Governs the controlled vocabulary. Operators can request new tags; admins edit, deprecate, "
                    "manage synonyms, delete (with migration), and approve/reject pending requests. The "
                    "tag-category taxonomy is admin-editable and importable from JSON. Tag actions cascade to "
                    "the chunks/documents that carry them.",
        "reach": "Tags tab.",
        "extra_shots": [
            ("07-tag-detail.png", "Tag selected: definition, tier/status and usage counts."),
            ("08-request-tag-modal.png", "Request-new-tag modal."),
        ],
        "stories": ["US-TAG-01", "US-TAG-02", "US-TAG-03", "US-TAG-04"],
        "cases": [
            {"id": "TC-TAG-01", "title": "Request a new tag and see it pending",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Tags tab open.",
             "steps": ["Click Request new tag.", "Fill name, definition, domain, justification.", "Submit request.",
                       "Leave and return to Tags."],
             "expected": "The requested tag appears in the pending/requested section and persists across navigation."},
            {"id": "TC-TAG-02", "title": "Edit a tag definition (admin)",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Admin; an existing tag.",
             "steps": ["Select a tag card.", "Click Edit.", "Change the definition.", "Save."],
             "expected": "The new definition is shown on the card and survives reload."},
            {"id": "TC-TAG-03", "title": "Approve / reject a pending request (admin)",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Admin; a pending request.",
             "steps": ["Reject one request — confirm the button is disabled until a reason is entered.",
                       "Approve another (optionally with edits)."],
             "expected": "Rejected tags show a Rejected status; approved ones become active; both emit activity events."},
            {"id": "TC-TAG-04", "title": "Manage synonyms (admin)",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Admin; a tag selected.",
             "steps": ["Open the synonyms action.", "Add and remove aliases.", "Save."],
             "expected": "The alias list updates and persists."},
            {"id": "TC-TAG-05", "title": "Deprecate a tag",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Admin; an active tag.",
             "steps": ["Open the deprecate action.", "Confirm."],
             "expected": "The tag is flagged deprecated and is excluded from default retrieval surfaces."},
            {"id": "TC-TAG-06", "title": "Delete a tag with migration to a replacement",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Admin; a tag used by documents.",
             "steps": ["Open Delete.", "Choose Migrate and pick a target tag (required).", "Confirm."],
             "expected": "Deletion requires a migration target (or explicit untag); affected chunks/docs are "
                         "migrated; the operation is auditable."},
            {"id": "TC-TAG-07", "title": "Edit the category taxonomy (admin)",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Admin; domain editor available.",
             "steps": ["Open the domain editor.", "Add a domain (id, label, colour).",
                       "Try an invalid id and an empty label.", "Save."],
             "expected": "Validation blocks invalid ids/labels/colours and duplicates; valid taxonomy saves."},
            {"id": "TC-TAG-08", "title": "Import categories from JSON",
             "type": "Manual", "sev": "S3",
             "pre": "Admin; a categories JSON (use Download template as a base).",
             "steps": ["Download the template.", "Import a valid JSON.", "Import a malformed JSON."],
             "expected": "Valid JSON replaces the taxonomy; malformed JSON is rejected with a clear error."},
        ],
    },
    {
        "title": "6.6 Retrieval",
        "shot": "10-retrieval-answer.png",
        "caption": "Retrieval tab: answer with a sources panel and the query-parameter rail.",
        "overview": "The question-answering surface. The operator asks a question and receives a streamed "
                    "answer with a sources panel; a parameter rail tunes query mode, top-k, rerank, token budget, "
                    "history turns and tag/document filters. Conversation threads persist in the browser.",
        "reach": "Retrieval tab.",
        "extra_shots": [("09-retrieval-empty.png", "Empty thread with the full query-parameter rail.")],
        "stories": ["US-RET-01", "US-RET-02", "US-RET-03"],
        "cases": [
            {"id": "TC-RET-01", "title": "Ask a question and receive an answer + sources",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Retrieval tab open.",
             "steps": ["Click New.", "Type a question in the query input.", "Click Send."],
             "expected": "An answer renders and a Sources panel lists the returned sources with scores."},
            {"id": "TC-RET-02", "title": "Threads persist across reload",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "At least one thread with a message.",
             "steps": ["Create two threads with questions.", "Switch between them.", "Reload."],
             "expected": "Both threads and their histories survive the reload; switching shows the right history."},
            {"id": "TC-RET-03", "title": "Delete a thread",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "A thread exists.",
             "steps": ["Delete a thread via its control."],
             "expected": "The thread is removed from the list and from local storage."},
            {"id": "TC-RET-04", "title": "Parameters change retrieval breadth",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Retrieval tab open.",
             "steps": ["Set top-k to 1.", "Ask a question."],
             "expected": "Fewer sources are returned than with a higher top-k (contract-level check on the mock)."},
            {"id": "TC-RET-05", "title": "Tag filter does NOT narrow the answer (known limitation)",
             "type": "Manual", "sev": "S4",
             "pre": "Retrieval tab; a tag selected in Source-tag filters.",
             "steps": ["Add a tag filter.", "Ask a question that has answers outside that tag."],
             "expected": "The answer may include out-of-tag material — this is expected (§4 item 1). Only file "
                         "a defect if the UI claims the filter is applied to the answer.",
             "note": "Do NOT report the no-op itself. See §4 / audit C1."},
            {"id": "TC-RET-06", "title": "Citation link drills to the document",
             "type": "E2E (to build)", "sev": "S3",
             "pre": "An answer with an inline citation marker.",
             "steps": ["Click an inline [n] citation or a source row."],
             "expected": "Navigates to / highlights the corresponding source document.",
             "note": "Currently skipped in automation; recommended new E2E and a manual pass."},
            {"id": "TC-RET-07", "title": "Answer quality and honest no-answer (real backend)",
             "type": "Manual", "sev": "S2",
             "pre": "Real backend with a populated KB.",
             "steps": ["Ask 10–15 in-domain questions.", "Ask 3 deliberately out-of-scope questions."],
             "expected": "In-domain answers are grounded with relevant sources; out-of-scope questions return "
                         "\"insufficient information\" with no fabricated grounding.",
             "note": "Mocked build returns canned text — judge quality on real backend only."},
        ],
    },
    {
        "title": "6.7 Knowledge graph",
        "shot": "12-graph-inspector.png",
        "caption": "Graph tab: canvas with an entity selected and its inspector open.",
        "overview": "A live projection of the knowledge graph. Operators filter by entity type, tag and "
                    "document, select nodes/edges to inspect, and pin nodes. Admins create, edit and delete "
                    "entities and relations; mutations emit activity and refetch the graph.",
        "reach": "Graph tab.",
        "extra_shots": [
            ("11-graph-canvas.png", "Graph canvas with the type/tag/document filter rail."),
            ("13-graph-add-entity.png", "Admin add-entity form."),
        ],
        "stories": ["US-GRAPH-01", "US-GRAPH-02", "US-GRAPH-03"],
        "cases": [
            {"id": "TC-GR-01", "title": "Filter by entity type updates canvas and counts",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Graph tab open.",
             "steps": ["Toggle entity-type filters.", "Observe nodes and the per-type counters."],
             "expected": "The canvas and counters track the active type filters."},
            {"id": "TC-GR-02", "title": "Search an entity by name",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Graph tab open.",
             "steps": ["Type an entity name in search."],
             "expected": "Only matching nodes remain emphasised/visible."},
            {"id": "TC-GR-03", "title": "Select an entity and read its inspector",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Graph has nodes.",
             "steps": ["Click a node.", "Read summary, type, tags, source documents and relations."],
             "expected": "The inspector shows the entity's details and its incoming/outgoing relations."},
            {"id": "TC-GR-04", "title": "Pin an entity; pin persists",
             "type": "E2E (automated)", "sev": "S4",
             "pre": "A node selected.",
             "steps": ["Click Pin.", "Reload and reselect the node."],
             "expected": "The node is pinned and remains pinned after reload."},
            {"id": "TC-GR-05", "title": "Create an entity (admin)",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Admin; Graph tab.",
             "steps": ["Click Add entity.", "Confirm submit is disabled until name + type are set.",
                       "Fill required fields and create."],
             "expected": "A new entity is created, appears on the canvas, and an activity event is recorded."},
            {"id": "TC-GR-06", "title": "Create a relation with duplicate guard (admin)",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Admin; an entity selected.",
             "steps": ["Add a relation to another entity.", "Try to add the same edge again."],
             "expected": "The relation is created; a duplicate edge is prevented."},
            {"id": "TC-GR-07", "title": "Delete an entity / relation (admin, two-click)",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Admin; an entity or relation selected.",
             "steps": ["Click Delete.", "Confirm on the second click."],
             "expected": "Deletion requires the second confirm; the element disappears; an event is recorded."},
            {"id": "TC-GR-08", "title": "Drill from an entity to its documents",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "An entity with source documents.",
             "steps": ["Use \"View documents mentioning this entity\"."],
             "expected": "Navigates to Documents pre-filtered by that entity."},
        ],
    },
    {
        "title": "6.8 Activity",
        "shot": "15-activity-detail.png",
        "caption": "Activity tab: filtered timeline with an event detail panel.",
        "overview": "The audit feed. Every governance action (upload, approve/reject, tag change, graph "
                    "mutation, key minting, classification rejection…) is recorded. Filter by time range, kind, "
                    "severity, actor and free text; open an event to drill into its target. Admins can clear activity.",
        "reach": "Activity tab.",
        "stories": ["US-ACT-01", "US-ACT-02"],
        "cases": [
            {"id": "TC-ACT-01", "title": "Severity filter narrows the timeline",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Events of mixed severity.",
             "steps": ["Set Severity filter to error.", "Reset to any."],
             "expected": "Only error rows show under the error filter; all return on reset; the stats count tracks the filter."},
            {"id": "TC-ACT-02", "title": "Kind and text filters",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Events of mixed kinds.",
             "steps": ["Toggle one or more event kinds.", "Search a text fragment."],
             "expected": "The list reflects the combined kind + text filters."},
            {"id": "TC-ACT-03", "title": "Time-range tabs",
             "type": "E2E (automated)", "sev": "S4",
             "pre": "Activity tab open.",
             "steps": ["Switch among 24h / 7d / 30d / All."],
             "expected": "The active range is highlighted and the timeline updates."},
            {"id": "TC-ACT-04", "title": "Open an event detail",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "At least one event.",
             "steps": ["Click an event row.", "Read the detail (Event ID, severity, actor, target, timestamp)."],
             "expected": "The detail panel shows full metadata for the selected event."},
            {"id": "TC-ACT-05", "title": "Drill from an event to its source",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "A source-related event selected.",
             "steps": ["In the detail panel, click Open source (or Replay ingestion for a failed source)."],
             "expected": "Navigates to the Documents tab on the right document (or replays ingestion)."},
            {"id": "TC-ACT-06", "title": "Refresh pulls new events",
             "type": "E2E (automated)", "sev": "S4",
             "pre": "Activity tab open.",
             "steps": ["Click Refresh after a new action elsewhere."],
             "expected": "Newly produced events appear."},
        ],
    },
    {
        "title": "6.9 Settings — profile, API, API keys, folders",
        "shot": "18-settings-api-keys.png",
        "caption": "Settings › API keys: mint, one-time reveal, and revoke per-operator keys.",
        "overview": "Four sections. Profile shows identity, IdP and gateway scopes (read-only). API browses the "
                    "live OpenAPI surface. API keys lets admins mint a key (full value revealed exactly once), "
                    "list prefixes, and revoke. Folders shows the env-controlled identity and, for admins, "
                    "runtime folder CRUD within a five-folder cap.",
        "reach": "Settings tab; left-rail section switch.",
        "extra_shots": [
            ("16-settings-profile.png", "Profile: identity, IdP and gateway scopes (read-only)."),
            ("17-settings-api.png", "Live OpenAPI surface browser."),
            ("19-settings-folders.png", "Runtime folder administration (admin:folders)."),
        ],
        "stories": ["US-SET-01", "US-SET-02", "US-SET-03", "US-SET-04"],
        "cases": [
            {"id": "TC-SET-01", "title": "Profile shows identity and scopes",
             "type": "E2E (automated)", "sev": "S4",
             "pre": "Signed in.",
             "steps": ["Open Settings › Profile."],
             "expected": "Name, email, role, IdP/realm/subject and gateway-scope chips are shown; an open-access "
                         "note appears when no auth backend is configured."},
            {"id": "TC-SET-02", "title": "API explorer loads and authorises",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Settings open.",
             "steps": ["Open the API rail.", "Authorize with a token, then revoke it."],
             "expected": "The OpenAPI surface lists endpoints; Authorize toggles to Authorized and back on revoke."},
            {"id": "TC-SET-03", "title": "Mint an API key with one-time reveal (admin)",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Admin; API keys rail.",
             "steps": ["Create a key with a name.", "Copy the revealed full value.", "Dismiss the reveal."],
             "expected": "The full value is shown once with a copy control and a \"never shown again\" warning; "
                         "afterwards only the prefix is listed."},
            {"id": "TC-SET-04", "title": "Revoke an API key (admin, two-click)",
             "type": "E2E (automated)", "sev": "S2",
             "pre": "Admin; an active key.",
             "steps": ["Click Revoke.", "Confirm on the second click."],
             "expected": "The key becomes revoked (kept for audit) and can no longer authenticate."},
            {"id": "TC-SET-05", "title": "Folder identity is read-only",
             "type": "Manual", "sev": "S4",
             "pre": "Settings › Folder.",
             "steps": ["Open the Folder section."],
             "expected": "Folder id and display name are shown as env-controlled and cannot be edited here."},
            {"id": "TC-SET-06", "title": "Create a runtime folder (admin)",
             "type": "E2E (automated)", "sev": "S3",
             "pre": "Admin with admin:folders; under the 5-folder cap.",
             "steps": ["Click Add folder.", "Enter a valid id, name and visibility.", "Create."],
             "expected": "The folder is created and appears in the topbar switcher."},
            {"id": "TC-SET-07", "title": "Folder cap and id validation",
             "type": "Manual", "sev": "S3",
             "pre": "Admin.",
             "steps": ["Try to create a 6th folder.", "Try an invalid id and a duplicate id."],
             "expected": "Add is blocked at the cap with an at-max indicator; invalid/duplicate ids are rejected with a clear error."},
            {"id": "TC-SET-08", "title": "Edit / delete a runtime folder; env folder is protected",
             "type": "Manual", "sev": "S2",
             "pre": "Admin; one env-seeded and one runtime folder.",
             "steps": ["Edit the runtime folder.", "Delete an empty runtime folder (two-click).",
                       "Attempt to edit/delete the env-seeded folder.", "Attempt to delete a folder that still has data."],
             "expected": "Runtime folders edit/delete; the env-seeded folder offers no edit/delete; deleting a "
                         "folder with data is refused (409)."},
            {"id": "TC-SET-09", "title": "Folder admin hidden for non-admins",
             "type": "Manual", "sev": "S1",
             "pre": "Reader or Steward without admin:folders (real/IdP config).",
             "steps": ["Open Settings › Folder."],
             "expected": "Folder CRUD is read-only / hidden; the API rejects folder mutations with 403.",
             "note": "Access-control case — a failure here is a security defect (S1)."},
        ],
    },
]


if __name__ == "__main__":
    build()
